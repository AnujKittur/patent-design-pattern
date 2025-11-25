from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_patent_llm_document():
    """Create Word document for Patent Documents & LLM for Design Research."""
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Research Proposal: LLM-Based Patent Document Analysis for Automatic Design Applications in Construction Robotics', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Individual Study Program - PhD Level Research')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.italic = True
    subtitle_format.size = Pt(14)
    
    doc.add_paragraph()
    
    # Table of Contents
    toc_heading = doc.add_heading('Table of Contents', 1)
    doc.add_paragraph('1. Executive Summary')
    doc.add_paragraph('2. Research Overview')
    doc.add_paragraph('3. Research Objectives')
    doc.add_paragraph('4. Theoretical Framework')
    doc.add_paragraph('5. Research Methodology')
    doc.add_paragraph('6. Data Sources & Repositories')
    doc.add_paragraph('7. Technical Implementation')
    doc.add_paragraph('8. Research Flowchart')
    doc.add_paragraph('   8.1 Component Interaction Diagram')
    doc.add_paragraph('   8.2 Detailed Workflows & Process Diagrams')
    doc.add_paragraph('       8.2.1 LLM Model Development Workflow')
    doc.add_paragraph('       8.2.2 Patent Document Processing Workflow')
    doc.add_paragraph('       8.2.3 Pattern Recognition & Data Mining Workflow')
    doc.add_paragraph('       8.2.4 Automatic Design Generation Workflow')
    doc.add_paragraph('       8.2.5 Data Flow Diagram')
    doc.add_paragraph('       8.2.6 System Architecture Flowchart')
    doc.add_paragraph('9. Timeline & Schedule')
    doc.add_paragraph('10. Expected Outcomes')
    doc.add_paragraph('11. References')
    
    doc.add_page_break()
    
    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', 1)
    doc.add_paragraph(
        'This research proposal focuses on leveraging Large Language Models (LLMs) to process vast repositories of '
        'patent documents and construction robotics technologies data for automatic design applications. The research '
        'will develop intelligent systems capable of extracting design patterns, specificities, and technical '
        'knowledge from textual and numeric data in patent documents. Using advanced data mining theories and LLM '
        'technologies, this research aims to build automated design systems that can generate optimized construction '
        'robotics designs based on patterns learned from patent databases. This PhD-level individual study program '
        '(2-3 credits) addresses the challenge of processing massive amounts of patent data and transforming it into '
        'actionable design knowledge for construction robotics applications.'
    )
    
    # 2. Research Overview
    doc.add_heading('2. Research Overview', 1)
    
    doc.add_heading('2.1 Core Research Topic', 2)
    doc.add_paragraph(
        'The primary research topic focuses on using Large Language Models (LLMs) to analyze patent documents and '
        'construction robotics technologies data for automatic design applications. Key components include:'
    )
    
    bullet_points = [
        'Processing vast repositories of patent documents using LLMs',
        'Extracting textual and numeric data from construction robotics technologies',
        'Data mining and pattern recognition from massive databases',
        'Building LLM models for successful data transfer and processing',
        'Automatic design application development for construction robotics'
    ]
    
    for point in bullet_points:
        p = doc.add_paragraph(point, style='List Bullet')
    
    doc.add_heading('2.2 Research Motivation', 2)
    doc.add_paragraph(
        'Vast repositories of data exist in patent documents and construction robotics technologies, containing '
        'valuable design knowledge, technical specifications, and innovation patterns. However, manually processing '
        'this massive amount of data is impractical. Large Language Models offer unprecedented capabilities for '
        'understanding and extracting meaningful information from these textual repositories, enabling the automatic '
        'generation of design applications based on learned patterns.'
    )
    
    doc.add_heading('2.3 Research Scope', 2)
    doc.add_paragraph('The research scope encompasses:')
    
    scope_items = [
        'Building and training LLM models for processing patent document text',
        'Extracting design patterns and specificities from patent databases',
        'Processing both textual and numeric data from construction robotics technologies',
        'Developing data mining algorithms for pattern recognition',
        'Creating automatic design applications that leverage extracted knowledge',
        'Building pattern databases enriched with construction robotics knowledge'
    ]
    
    for item in scope_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 3. Research Objectives
    doc.add_heading('3. Research Objectives', 1)
    
    objectives = [
        {
            'title': 'Primary Objective 1: LLM Model Development for Patent Processing',
            'description': 'Develop and fine-tune Large Language Models specifically designed for processing patent '
                          'documents related to construction robotics technologies. The models should be capable of '
                          'understanding technical terminology, extracting design specifications, and identifying '
                          'innovative patterns.'
        },
        {
            'title': 'Primary Objective 2: Textual and Numeric Data Extraction',
            'description': 'Build systems capable of extracting both textual information (descriptions, claims, '
                          'abstracts) and numeric data (dimensions, performance metrics, specifications) from patent '
                          'documents and construction robotics datasets. Ensure successful data transfer and '
                          'normalization.'
        },
        {
            'title': 'Primary Objective 3: Pattern Recognition & Data Mining',
            'description': 'Apply advanced data mining theories to identify recurring patterns, design specificities, '
                          'and knowledge structures from vast patent repositories. Build pattern databases that '
                          'capture design principles and technical knowledge.'
        },
        {
            'title': 'Primary Objective 4: Automatic Design Application Development',
            'description': 'Develop intelligent systems for automatic design applications in construction robotics, '
                          'leveraging patterns and knowledge extracted from patent documents. The system should '
                          'generate optimized designs based on learned patterns and design specificities.'
        },
        {
            'title': 'Primary Objective 5: Design Specificity Identification',
            'description': 'Identify and categorize design specificities and constraints from patent documents, '
                          'creating a comprehensive knowledge base that informs automatic design generation.'
        }
    ]
    
    for i, obj in enumerate(objectives, 1):
        doc.add_heading(f'3.{i} {obj["title"]}', 2)
        doc.add_paragraph(obj['description'])
    
    # 4. Theoretical Framework
    doc.add_heading('4. Theoretical Framework', 1)
    
    doc.add_heading('4.1 Large Language Models (LLMs)', 2)
    doc.add_paragraph(
        'Large Language Models represent the state-of-the-art in natural language understanding and generation. '
        'For this research, LLMs will be used to:'
    )
    
    llm_aspects = [
        'Understand technical patent language and terminology',
        'Extract structured information from unstructured patent text',
        'Identify relationships between patent concepts',
        'Generate summaries and insights from patent documents',
        'Transfer knowledge learned from patents to design applications'
    ]
    
    for aspect in llm_aspects:
        doc.add_paragraph(aspect, style='List Bullet')
    
    doc.add_heading('4.2 Data Mining Theory', 2)
    doc.add_paragraph(
        'Data mining theories provide the foundation for extracting patterns and knowledge from large datasets. '
        'Key approaches include:'
    )
    
    mining_aspects = [
        'Pattern recognition algorithms for identifying design patterns',
        'Association rule mining for discovering relationships',
        'Clustering techniques for grouping similar designs',
        'Classification methods for categorizing design elements',
        'Knowledge graph construction for representing relationships'
    ]
    
    for aspect in mining_aspects:
        doc.add_paragraph(aspect, style='List Bullet')
    
    doc.add_heading('4.3 Automatic Design Theory', 2)
    doc.add_paragraph(
        'Automatic design applications leverage computational methods to generate designs based on constraints, '
        'objectives, and learned patterns. This research will integrate:'
    )
    
    design_aspects = [
        'Generative design methodologies',
        'Constraint-based design systems',
        'Pattern-based design generation',
        'Optimization algorithms for design refinement',
        'Knowledge-guided design synthesis'
    ]
    
    for aspect in design_aspects:
        doc.add_paragraph(aspect, style='List Bullet')
    
    # 5. Research Methodology
    doc.add_heading('5. Research Methodology', 1)
    
    doc.add_heading('5.1 Data Collection Phase', 2)
    doc.add_paragraph('Phase 1: Collection of vast repositories of data from:')
    collection_items = [
        'Patent documents databases (USPTO, EPO, WIPO)',
        'Construction robotics technologies datasets',
        'Technical specifications and design documents',
        'Performance metrics and numeric data',
        'Case studies and application examples'
    ]
    for item in collection_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('5.2 Data Preprocessing Phase', 2)
    doc.add_paragraph('Phase 2: Preprocessing collected data:')
    preprocessing_items = [
        'Text normalization and cleaning',
        'Numeric data extraction and standardization',
        'Format conversion and standardization',
        'Data validation and quality assessment',
        'Database construction and indexing'
    ]
    for item in preprocessing_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('5.3 LLM Model Development Phase', 2)
    doc.add_paragraph('Phase 3: Building and training LLM models:')
    llm_items = [
        'Model architecture selection and design',
        'Domain-specific fine-tuning on patent documents',
        'Text extraction and understanding training',
        'Multi-modal processing for text and numeric data',
        'Pattern recognition capability development'
    ]
    for item in llm_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('5.4 Data Mining & Pattern Recognition Phase', 2)
    doc.add_paragraph('Phase 4: Pattern extraction and knowledge discovery:')
    mining_items = [
        'Pattern identification algorithms implementation',
        'Design specificity extraction',
        'Knowledge graph construction',
        'Pattern database creation and enrichment',
        'Relationship mapping between design elements'
    ]
    for item in mining_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('5.5 Automatic Design Application Phase', 2)
    doc.add_paragraph('Phase 5: Building automatic design systems:')
    design_items = [
        'Design generation algorithm development',
        'Constraint-based design synthesis',
        'Pattern integration and application',
        'Optimization and refinement processes',
        'Validation and quality assessment mechanisms'
    ]
    for item in design_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 6. Data Sources & Repositories
    doc.add_heading('6. Data Sources & Repositories', 1)
    
    doc.add_heading('6.1 Patent Documents', 2)
    doc.add_paragraph(
        'Vast repositories of patent documents will serve as primary data sources. These documents contain rich '
        'textual information about construction technologies, robotics innovations, and technical specifications. '
        'Key sources include:'
    )
    patent_sources = [
        'United States Patent and Trademark Office (USPTO) - Full-text patents',
        'European Patent Office (EPO) - European patent database',
        'World Intellectual Property Organization (WIPO) - Global patent database',
        'Google Patents - Comprehensive patent search database',
        'IEEE Xplore - Technical patents and standards',
        'Patent databases specific to construction and robotics'
    ]
    for source in patent_sources:
        doc.add_paragraph(source, style='List Bullet')
    
    doc.add_heading('6.2 Construction Robotics Technologies', 2)
    doc.add_paragraph(
        'Datasets containing both textual and numeric data related to construction robotics technologies will be '
        'collected. This includes:'
    )
    robotics_data = [
        'Technical specifications and design documents',
        'Performance metrics and numeric data (dimensions, capacities, speeds)',
        'Application case studies and implementation reports',
        'Automation protocols and standards',
        'Design patterns and methodologies',
        'Component libraries and specifications'
    ]
    for data in robotics_data:
        doc.add_paragraph(data, style='List Bullet')
    
    doc.add_heading('6.3 Pattern Databases', 2)
    doc.add_paragraph(
        'Existing pattern databases will be utilized and enhanced through data mining theories to identify '
        'recurring patterns, design specificities, and optimal configurations. New pattern databases will be '
        'constructed specifically for construction robotics applications.'
    )
    
    # 7. Technical Implementation
    doc.add_heading('7. Technical Implementation', 1)
    
    doc.add_heading('7.1 LLM Architecture & Training', 2)
    doc.add_paragraph(
        'Large Language Models will be developed and fine-tuned specifically for patent document processing:'
    )
    llm_impl = [
        'Base model selection (GPT, BERT, T5, or specialized architectures)',
        'Domain-specific fine-tuning on patent corpora',
        'Multi-task learning for various extraction tasks',
        'Few-shot and zero-shot learning capabilities',
        'Efficient inference optimization for large-scale processing'
    ]
    for item in llm_impl:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('7.2 Data Extraction Pipeline', 2)
    doc.add_paragraph(
        'Automated pipelines for extracting information from patent documents:'
    )
    extraction_impl = [
        'Text extraction and normalization from PDF/XML formats',
        'Structured data extraction (claims, specifications, drawings)',
        'Numeric data extraction and validation',
        'Entity recognition for technical terms and concepts',
        'Relationship extraction between design elements'
    ]
    for item in extraction_impl:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('7.3 Pattern Recognition System', 2)
    doc.add_paragraph(
        'Systems for identifying and classifying patterns:'
    )
    pattern_impl = [
        'Supervised learning for pattern classification',
        'Unsupervised learning for pattern discovery',
        'Knowledge graph construction and querying',
        'Similarity matching and clustering algorithms',
        'Pattern database management and retrieval'
    ]
    for item in pattern_impl:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('7.4 Automatic Design Generation', 2)
    doc.add_paragraph(
        'Systems for generating designs automatically:'
    )
    generation_impl = [
        'Generative models for design synthesis',
        'Constraint satisfaction algorithms',
        'Multi-objective optimization for design quality',
        'Pattern-based design composition',
        'Design validation and feasibility checking'
    ]
    for item in generation_impl:
        doc.add_paragraph(item, style='List Bullet')
    
    # 8. Research Flowchart
    doc.add_heading('8. Research Flowchart', 1)
    doc.add_paragraph('The following flowchart illustrates the complete research workflow:')
    
    flowchart_text = """
    ┌─────────────────────────────────────────────────────────────────────────┐
    │        PATENT DOCUMENTS & LLM FOR DESIGN RESEARCH WORKFLOW              │
    └─────────────────────────────────────────────────────────────────────────┘
    
    START
     │
     ├─► [1. PATENT DATA COLLECTION]
     │    │
     │    ├─► USPTO Database
     │    ├─► EPO Database
     │    ├─► WIPO Database
     │    └─► Construction Robotics Patents
     │
     ├─► [2. CONSTRUCTION ROBOTICS DATA COLLECTION]
     │    │
     │    ├─► Textual Data Collection
     │    ├─► Numeric Data Collection
     │    ├─► Design Documents
     │    └─► Technical Specifications
     │
     ├─► [3. DATA PREPROCESSING]
     │    │
     │    ├─► Text Normalization
     │    ├─► Data Cleaning
     │    ├─► Format Standardization
     │    ├─► Numeric Data Extraction
     │    └─► Quality Validation
     │
     ├─► [4. LLM MODEL DEVELOPMENT]
     │    │
     │    ├─► Base Model Selection
     │    ├─► Domain Fine-tuning
     │    ├─► Text Understanding Training
     │    ├─► Multi-modal Integration
     │    └─► Pattern Recognition Training
     │
     ├─► [5. TEXTUAL DATA PROCESSING]
     │    │
     │    ├─► Patent Document Parsing
     │    ├─► Technical Term Extraction
     │    ├─► Claim Analysis
     │    ├─► Description Understanding
     │    └─► Relationship Extraction
     │
     ├─► [6. NUMERIC DATA PROCESSING]
     │    │
     │    ├─► Specification Extraction
     │    ├─► Performance Metric Collection
     │    ├─► Dimensional Data Processing
     │    └─► Data Standardization
     │
     ├─► [7. DATA MINING & PATTERN RECOGNITION]
     │    │
     │    ├─► Pattern Identification
     │    ├─► Design Specificity Extraction
     │    ├─► Knowledge Graph Construction
     │    ├─► Association Rule Mining
     │    └─► Pattern Database Creation
     │
     ├─► [8. DESIGN SPECIFICITY ANALYSIS]
     │    │
     │    ├─► Specificity Identification
     │    ├─► Constraint Extraction
     │    ├─► Design Principle Discovery
     │    └─► Categorization & Classification
     │
     ├─► [9. AUTOMATIC DESIGN GENERATION]
     │    │
     │    ├─► Pattern-Based Generation
     │    ├─► Constraint Integration
     │    ├─► Design Synthesis
     │    ├─► Optimization Processes
     │    └─► Quality Validation
     │
     ├─► [10. PATTERN DATABASE ENRICHMENT]
     │    │
     │    ├─► Knowledge Integration
     │    ├─► Pattern Refinement
     │    ├─► Relationship Mapping
     │    └─► Database Updates
     │
     ├─► [11. SYSTEM INTEGRATION]
     │    │
     │    ├─► Component Integration
     │    ├─► Workflow Orchestration
     │    ├─► API Development
     │    └─► User Interface Creation
     │
     └─► [12. TESTING & VALIDATION]
          │
          ├─► Design Quality Assessment
          ├─► Pattern Accuracy Testing
          ├─► System Performance Evaluation
          ├─► Case Study Validation
          └─► Documentation & Publication
          
    END
    """
    
    flowchart_para = doc.add_paragraph()
    flowchart_para.add_run(flowchart_text).font.name = 'Courier New'
    flowchart_para.add_run().font.size = Pt(9)
    
    doc.add_paragraph()
    
    # Component interaction diagram
    doc.add_heading('8.1 Component Interaction Diagram', 2)
    interaction_text = """
    ┌─────────────────────────────────────────────────────────────────────────┐
    │                    COMPONENT INTERACTION DIAGRAM                        │
    └─────────────────────────────────────────────────────────────────────────┘
    
    Patent Documents ──► LLM Processing ──► Text Extraction ──► Pattern Recognition
         │                    │                    │                    │
         │                    │                    │                    │
         ▼                    ▼                    ▼                    ▼
    Construction    ──► Numeric Data ──► Data Mining ──► Pattern Database
    Robotics Data      Extraction       Algorithms       │
         │                    │                    │                    │
         └────────────────────┼────────────────────┼────────────────────┘
                               │                    │
                               ▼                    ▼
                        Design Specificities ──► Knowledge Base
                               │                    │
                               └────────────────────┘
                                         │
                                         ▼
                            Automatic Design Generator
                                         │
                                         ▼
                              Optimized Construction
                              Robotics Designs
    """
    
    interaction_para = doc.add_paragraph()
    interaction_para.add_run(interaction_text).font.name = 'Courier New'
    interaction_para.add_run().font.size = Pt(9)
    
    doc.add_page_break()
    
    # Detailed Workflows Section
    doc.add_heading('8.2 Detailed Workflows & Process Diagrams', 1)
    
    doc.add_heading('8.2.1 LLM Model Development Workflow', 2)
    doc.add_paragraph('Detailed workflow for LLM model development and fine-tuning:')
    
    llm_workflow = """
    ┌─────────────────────────────────────────────────────────────────────────┐
    │              LLM MODEL DEVELOPMENT WORKFLOW                             │
    └─────────────────────────────────────────────────────────────────────────┘
    
    START LLM Development
     │
     ├─► [STEP 1: Base Model Selection]
     │    │
     │    ├─► Evaluate Model Options
     │    │    ├─► GPT Models
     │    │    ├─► BERT Models
     │    │    ├─► T5 Models
     │    │    └─► Specialized Architectures
     │    │
     │    ├─► Assess Model Capabilities
     │    ├─► Compare Model Sizes
     │    └─► Select Optimal Base Model
     │         │
     │         └─► Output: Selected Base Model
     │
     ├─► [STEP 2: Patent Corpus Preparation]
     │    │
     │    ├─► Collect Patent Documents
     │    ├─► Extract Text Content
     │    ├─► Clean and Normalize Text
     │    ├─► Create Training Dataset
     │    └─► Partition Train/Val/Test Sets
     │         │
     │         └─► Output: Prepared Corpus
     │
     ├─► [STEP 3: Domain-Specific Fine-tuning]
     │    │
     │    ├─► Task Definition
     │    │    ├─► Text Classification
     │    │    ├─► Named Entity Recognition
     │    │    ├─► Information Extraction
     │    │    └─► Pattern Recognition
     │    │
     │    ├─► Fine-tuning Configuration
     │    ├─► Training Execution
     │    ├─► Hyperparameter Tuning
     │    └─► Model Checkpointing
     │         │
     │         └─► Output: Fine-tuned Model
     │
     ├─► [STEP 4: Model Evaluation]
     │    │
     │    ├─► Performance Metrics
     │    │    ├─► Accuracy
     │    │    ├─► Precision/Recall
     │    │    ├─► F1-Score
     │    │    └─► Domain-Specific Metrics
     │    │
     │    ├─► Test Set Evaluation
     │    ├─► Error Analysis
     │    └─► Model Comparison
     │         │
     │         └─► Output: Evaluation Report
     │
     ├─► [STEP 5: Model Optimization]
     │    │
     │    ├─► Identify Improvement Areas
     │    ├─► Adjust Fine-tuning Parameters
     │    ├─► Data Augmentation
     │    ├─► Iterative Refinement
     │    └─► Export Final Model
     │         │
     │         └─► Output: Optimized Model
     │
     └─► [STEP 6: Model Deployment]
          │
          ├─► Model Serialization
          ├─► API Development
          ├─► Inference Pipeline Setup
          └─► Performance Monitoring
               │
               └─► Output: Deployed Model
                    │
                    └─► END
    """
    
    llm_workflow_para = doc.add_paragraph()
    llm_workflow_para.add_run(llm_workflow).font.name = 'Courier New'
    llm_workflow_para.add_run().font.size = Pt(8)
    
    doc.add_paragraph()
    
    doc.add_heading('8.2.2 Patent Document Processing Workflow', 2)
    doc.add_paragraph('Detailed workflow for processing patent documents:')
    
    patent_workflow = """
    ┌─────────────────────────────────────────────────────────────────────────┐
    │           PATENT DOCUMENT PROCESSING WORKFLOW                           │
    └─────────────────────────────────────────────────────────────────────────┘
    
    START Patent Processing
     │
     ├─► [STEP 1: Document Acquisition]
     │    │
     │    ├─► Query Patent Databases
     │    ├─► Filter by Construction/Robotics
     │    ├─► Download Patent Documents
     │    └─► Organize Documents
     │         │
     │         └─► Output: Raw Patent Collection
     │
     ├─► [STEP 2: Document Parsing]
     │    │
     │    ├─► Extract Text from PDF/XML
     │    ├─► Parse Document Structure
     │    │    ├─► Title
     │    │    ├─► Abstract
     │    │    ├─► Claims
     │    │    ├─► Description
     │    │    └─► Drawings/Figures
     │    │
     │    └─► Output: Structured Patent Data
     │
     ├─► [STEP 3: Textual Data Extraction]
     │    │
     │    ├─► LLM-Based Text Understanding
     │    ├─► Technical Term Extraction
     │    ├─► Concept Identification
     │    ├─► Relationship Extraction
     │    └─► Summary Generation
     │         │
     │         └─► Output: Extracted Text Features
     │
     ├─► [STEP 4: Numeric Data Extraction]
     │    │
     │    ├─► Specification Extraction
     │    ├─► Dimension Extraction
     │    ├─► Performance Metric Extraction
     │    ├─► Parameter Value Extraction
     │    └─► Data Normalization
     │         │
     │         └─► Output: Extracted Numeric Data
     │
     ├─► [STEP 5: Data Integration]
     │    │
     │    ├─► Merge Textual & Numeric Data
     │    ├─► Create Patent Records
     │    ├─► Link Related Patents
     │    └─► Build Patent Knowledge Base
     │         │
     │         └─► Output: Integrated Patent Database
     │
     ├─► [STEP 6: Quality Validation]
     │    │
     │    ├─► Data Completeness Check
     │    ├─► Data Accuracy Verification
     │    ├─► Consistency Validation
     │    └─► Error Correction
     │         │
     │         └─► Output: Validated Database
     │
     └─► [STEP 7: Database Storage]
          │
          ├─► Index Patent Records
          ├─► Create Searchable Index
          ├─► Implement Query Interface
          └─► Backup & Archive
               │
               └─► Output: Patent Database
                    │
                    └─► END
    """
    
    patent_workflow_para = doc.add_paragraph()
    patent_workflow_para.add_run(patent_workflow).font.name = 'Courier New'
    patent_workflow_para.add_run().font.size = Pt(8)
    
    doc.add_paragraph()
    
    doc.add_heading('8.2.3 Pattern Recognition & Data Mining Workflow', 2)
    doc.add_paragraph('Detailed workflow for pattern recognition and data mining:')
    
    pattern_workflow = """
    ┌─────────────────────────────────────────────────────────────────────────┐
    │        PATTERN RECOGNITION & DATA MINING WORKFLOW                      │
    └─────────────────────────────────────────────────────────────────────────┘
    
    START Pattern Mining
     │
     ├─► [STEP 1: Pattern Discovery]
     │    │
     │    ├─► Association Rule Mining
     │    │    ├─► Find Frequent Itemsets
     │    │    ├─► Generate Association Rules
     │    │    └─► Calculate Confidence/Support
     │    │
     │    ├─► Clustering Analysis
     │    │    ├─► Group Similar Designs
     │    │    ├─► Identify Design Families
     │    │    └─► Cluster Technical Features
     │    │
     │    └─► Output: Initial Patterns
     │
     ├─► [STEP 2: Design Specificity Extraction]
     │    │
     │    ├─► Extract Design Constraints
     │    ├─► Identify Design Requirements
     │    ├─► Capture Design Preferences
     │    ├─► Extract Technical Specifications
     │    └─► Categorize Specificities
     │         │
     │         └─► Output: Design Specificity Set
     │
     ├─► [STEP 3: Knowledge Graph Construction]
     │    │
     │    ├─► Extract Entities
     │    ├─► Identify Relationships
     │    ├─► Build Graph Structure
     │    ├─► Enrich with Metadata
     │    └─► Validate Graph Integrity
     │         │
     │         └─► Output: Knowledge Graph
     │
     ├─► [STEP 4: Pattern Validation]
     │    │
     │    ├─► Assess Pattern Significance
     │    ├─► Validate Pattern Frequency
     │    ├─► Verify Pattern Relevance
     │    └─► Filter Low-Quality Patterns
     │         │
     │         └─► Output: Validated Patterns
     │
     ├─► [STEP 5: Pattern Database Construction]
     │    │
     │    ├─► Structure Pattern Storage
     │    ├─► Index Patterns by Category
     │    ├─► Link Related Patterns
     │    ├─► Add Metadata & Annotations
     │    └─► Implement Pattern Query System
     │         │
     │         └─► Output: Pattern Database
     │
     └─► [STEP 6: Pattern Application]
          │
          ├─► Pattern Retrieval System
          ├─► Pattern Matching Algorithms
          ├─► Relevance Scoring
          └─► Pattern Recommendation
               │
               └─► Output: Pattern Application System
                    │
                    └─► END
    """
    
    pattern_workflow_para = doc.add_paragraph()
    pattern_workflow_para.add_run(pattern_workflow).font.name = 'Courier New'
    pattern_workflow_para.add_run().font.size = Pt(8)
    
    doc.add_paragraph()
    
    doc.add_heading('8.2.4 Automatic Design Generation Workflow', 2)
    doc.add_paragraph('Detailed workflow for automatic design generation:')
    
    design_workflow = """
    ┌─────────────────────────────────────────────────────────────────────────┐
    │           AUTOMATIC DESIGN GENERATION WORKFLOW                         │
    └─────────────────────────────────────────────────────────────────────────┘
    
    START Design Generation
     │
     ├─► [STEP 1: Requirements Input]
     │    │
     │    ├─► Accept User Requirements
     │    ├─► Parse Design Constraints
     │    ├─► Identify Design Objectives
     │    └─► Validate Input Completeness
     │         │
     │         └─► Output: Design Requirements
     │
     ├─► [STEP 2: Pattern Retrieval]
     │    │
     │    ├─► Query Pattern Database
     │    ├─► Match Requirements to Patterns
     │    ├─► Retrieve Relevant Patterns
     │    ├─► Rank Patterns by Relevance
     │    └─► Select Top Patterns
     │         │
     │         └─► Output: Selected Patterns
     │
     ├─► [STEP 3: Design Specificity Matching]
     │    │
     │    ├─► Extract Relevant Specificities
     │    ├─► Match Specificities to Requirements
     │    ├─► Identify Compatible Specificities
     │    └─► Resolve Specificity Conflicts
     │         │
     │         └─► Output: Design Specificity Set
     │
     ├─► [STEP 4: Initial Design Generation]
     │    │
     │    ├─► Pattern-Based Synthesis
     │    ├─► Constraint Application
     │    ├─► Component Integration
     │    └─► Generate Initial Design
     │         │
     │         └─► Output: Initial Design
     │
     ├─► [STEP 5: Design Optimization]
     │    │
     │    ├─► Multi-Objective Optimization
     │    │    ├─► Performance Optimization
     │    │    ├─► Cost Optimization
     │    │    └─► Reliability Optimization
     │    │
     │    ├─► Iterative Refinement
     │    ├─► Constraint Satisfaction
     │    └─► Quality Improvement
     │         │
     │         └─► Output: Optimized Design
     │
     ├─► [STEP 6: Design Validation]
     │    │
     │    ├─► Feasibility Check
     │    ├─► Constraint Verification
     │    ├─► Performance Evaluation
     │    ├─► Quality Assessment
     │    └─► Error Detection
     │         │
     │         └─► Output: Validation Report
     │
     ├─► [STEP 7: Design Refinement (if needed)]
     │    │
     │    ├─► Identify Issues
     │    ├─► Modify Design
     │    ├─► Re-optimize
     │    └─► Re-validate
     │         │
     │         └─► Loop back to Step 5 if needed
     │
     └─► [STEP 8: Final Design Output]
          │
          ├─► Generate Design Documentation
          ├─► Create Design Specifications
          ├─► Produce Design Visualizations
          └─► Export Design Files
               │
               └─► Output: Final Design
                    │
                    └─► END
    """
    
    design_workflow_para = doc.add_paragraph()
    design_workflow_para.add_run(design_workflow).font.name = 'Courier New'
    design_workflow_para.add_run().font.size = Pt(8)
    
    doc.add_paragraph()
    
    doc.add_heading('8.2.5 Data Flow Diagram', 2)
    doc.add_paragraph('Data flow through the entire system:')
    
    data_flow = """
    ┌─────────────────────────────────────────────────────────────────────────┐
    │                        DATA FLOW DIAGRAM                                │
    └─────────────────────────────────────────────────────────────────────────┘
    
    ┌────────────────────┐
    │  Patent Documents  │────────────────┐
    │  (USPTO/EPO/WIPO)  │                │
    └────────────────────┘                │
                                          │
    ┌────────────────────┐                │
    │  Construction       │────────────────┤
    │  Robotics Data     │                │
    │  (Text & Numeric)  │                │
    └────────────────────┘                │
                                          ▼
                          ┌───────────────────────────┐
                          │   Data Preprocessing      │
                          │   & Normalization         │
                          └───────────────────────────┘
                                          │
                                          ▼
                          ┌───────────────────────────┐
                          │   LLM Processing         │
                          │   Pipeline               │
                          └───────────────────────────┘
                                          │
                      ┌───────────────────┴───────────────────┐
                      │                                       │
                      ▼                                       ▼
      ┌───────────────────────────┐      ┌───────────────────────────┐
      │   Textual Data            │      │   Numeric Data            │
      │   Extraction              │      │   Extraction              │
      └───────────────────────────┘      └───────────────────────────┘
                      │                                       │
                      └───────────────┬───────────────────────┘
                                      │
                                      ▼
                      ┌───────────────────────────┐
                      │   Integrated Patent       │
                      │   Database               │
                      └───────────────────────────┘
                                      │
                                      ▼
                      ┌───────────────────────────┐
                      │   Data Mining &          │
                      │   Pattern Recognition    │
                      └───────────────────────────┘
                                      │
                      ┌───────────────┴───────────────┐
                      │                               │
                      ▼                               ▼
      ┌───────────────────────────┐      ┌───────────────────────────┐
      │   Pattern Database        │      │   Knowledge Graph          │
      └───────────────────────────┘      └───────────────────────────┘
                      │                               │
                      └───────────────┬───────────────┘
                                      │
                                      ▼
                      ┌───────────────────────────┐
                      │   Design Specificity     │
                      │   Knowledge Base         │
                      └───────────────────────────┘
                                      │
                                      ▼
                      ┌───────────────────────────┐
                      │   Automatic Design        │
                      │   Generator              │
                      └───────────────────────────┘
                                      │
                                      ▼
                      ┌───────────────────────────┐
                      │   Generated Designs       │
                      │   (Construction Robotics) │
                      └───────────────────────────┘
    """
    
    dataflow_para = doc.add_paragraph()
    dataflow_para.add_run(data_flow).font.name = 'Courier New'
    dataflow_para.add_run().font.size = Pt(8)
    
    doc.add_paragraph()
    
    doc.add_heading('8.2.6 System Architecture Flowchart', 2)
    doc.add_paragraph('Complete system architecture and component interactions:')
    
    architecture_flow = """
    ┌─────────────────────────────────────────────────────────────────────────┐
    │                    SYSTEM ARCHITECTURE FLOWCHART                        │
    └─────────────────────────────────────────────────────────────────────────┘
    
    ┌─────────────────────────────────────────────────────────────────────┐
    │                         UI/API LAYER                                │
    │  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐              │
    │  │ User         │  │ Design       │  │ Pattern      │              │
    │  │ Interface    │  │ Request API  │  │ Query API    │              │
    │  └──────────────┘  └──────────────┘  └──────────────┘              │
    └─────────────────────────────────────────────────────────────────────┘
                                        │
                                        ▼
    ┌─────────────────────────────────────────────────────────────────────┐
    │                      APPLICATION LAYER                              │
    │  ┌──────────────────────────────────────────────────────────────┐  │
    │  │          Automatic Design Generation Engine                    │  │
    │  │  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐       │  │
    │  │  │ Pattern       │  │ Design       │  │ Optimization │       │  │
    │  │  │ Matcher       │  │ Synthesizer  │  │ Engine       │       │  │
    │  │  └──────────────┘  └──────────────┘  └──────────────┘       │  │
    │  └──────────────────────────────────────────────────────────────┘  │
    └─────────────────────────────────────────────────────────────────────┘
                                        │
                    ┌───────────────────┴───────────────────┐
                    │                                       │
                    ▼                                       ▼
    ┌───────────────────────────┐      ┌───────────────────────────┐
    │   Pattern Recognition     │      │   LLM Processing          │
    │   & Data Mining Module    │      │   Module                  │
    │                           │      │                           │
    │  ┌──────────────────────┐  │      │  ┌──────────────────────┐ │
    │  │ Association Mining   │  │      │  │ Text Extraction       │ │
    │  │ Clustering           │  │      │  │ Entity Recognition   │ │
    │  │ Classification      │  │      │  │ Relationship Extract  │ │
    │  └──────────────────────┘  │      │  └──────────────────────┘ │
    └───────────────────────────┘      └───────────────────────────┘
                    │                                       │
                    └───────────────┬───────────────────────┘
                                    │
                                    ▼
    ┌─────────────────────────────────────────────────────────────────────┐
    │                         DATA LAYER                                   │
    │  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐              │
    │  │ Patent        │  │ Pattern      │  │ Knowledge    │              │
    │  │ Database      │  │ Database     │  │ Graph DB     │              │
    │  └──────────────┘  └──────────────┘  └──────────────┘              │
    └─────────────────────────────────────────────────────────────────────┘
                                        │
                                        ▼
    ┌─────────────────────────────────────────────────────────────────────┐
    │                      EXTERNAL SOURCES                                │
    │  ┌──────────────┐  ┌──────────────┐  ┌──────────────┐              │
    │  │ USPTO        │  │ EPO          │  │ WIPO         │              │
    │  │ Database     │  │ Database     │  │ Database     │              │
    │  └──────────────┘  └──────────────┘  └──────────────┘              │
    └─────────────────────────────────────────────────────────────────────┘
    """
    
    arch_para = doc.add_paragraph()
    arch_para.add_run(architecture_flow).font.name = 'Courier New'
    arch_para.add_run().font.size = Pt(8)
    
    # 9. Timeline & Schedule
    doc.add_heading('9. Timeline & Schedule', 1)
    
    doc.add_heading('9.1 Meeting Schedule', 2)
    doc.add_paragraph('Regular meetings will be conducted:')
    schedule_items = [
        'Days: Monday through Thursday',
        'Time Options: 5:00 PM - 7:30 PM OR 7:00 PM - 9:30 PM',
        'Frequency: Weekly or as determined by research progress',
        'Format: Individual study program supervision'
    ]
    for item in schedule_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('9.2 Research Phases Timeline', 2)
    
    timeline_data = [
        {'phase': 'Phase 1: Literature Review & Framework Development', 'duration': 'Weeks 1-4'},
        {'phase': 'Phase 2: Data Collection & Preprocessing', 'duration': 'Weeks 5-12'},
        {'phase': 'Phase 3: LLM Model Development & Training', 'duration': 'Weeks 13-22'},
        {'phase': 'Phase 4: Textual & Numeric Data Extraction', 'duration': 'Weeks 23-28'},
        {'phase': 'Phase 5: Data Mining & Pattern Recognition', 'duration': 'Weeks 29-36'},
        {'phase': 'Phase 6: Design Specificity Identification', 'duration': 'Weeks 37-40'},
        {'phase': 'Phase 7: Automatic Design Application Development', 'duration': 'Weeks 41-48'},
        {'phase': 'Phase 8: Pattern Database Construction', 'duration': 'Weeks 49-50'},
        {'phase': 'Phase 9: Integration, Testing & Validation', 'duration': 'Weeks 51-52'}
    ]
    
    for timeline in timeline_data:
        p = doc.add_paragraph()
        p.add_run(f'{timeline["phase"]}: ').bold = True
        p.add_run(timeline['duration'])
    
    doc.add_heading('9.3 Credit Allocation', 2)
    doc.add_paragraph(
        'This individual study program is allocated 2-3 credit hours, reflecting the intensive nature of '
        'PhD-level research work.'
    )
    
    # 10. Expected Outcomes
    doc.add_heading('10. Expected Outcomes', 1)
    
    doc.add_heading('10.1 Technical Contributions', 2)
    technical_outcomes = [
        'Advanced LLM models fine-tuned for patent document processing',
        'Novel data extraction pipelines for patent documents',
        'Data mining algorithms optimized for construction robotics patterns',
        'Automatic design generation systems',
        'Comprehensive pattern databases for construction robotics'
    ]
    for outcome in technical_outcomes:
        doc.add_paragraph(outcome, style='List Bullet')
    
    doc.add_heading('10.2 Practical Applications', 2)
    practical_outcomes = [
        'Automated design tools for construction robotics',
        'Pattern recognition systems for patent analysis',
        'Knowledge extraction systems from patent repositories',
        'Design recommendation engines',
        'Innovation discovery tools for construction technologies'
    ]
    for outcome in practical_outcomes:
        doc.add_paragraph(outcome, style='List Bullet')
    
    doc.add_heading('10.3 Research Deliverables', 2)
    doc.add_paragraph(
        'The primary deliverable of this research project is a fully functional RAG (Retrieval-Augmented Generation) '
        'based automatic design generation system for construction robotics. The system will enable users to input '
        'design specifications as prompts and receive optimized design solutions generated from patent document knowledge. '
        'The following deliverables will be provided:'
    )
    doc.add_paragraph()
    
    doc.add_heading('10.3.1 Core System Deliverable: RAG-Based Design Generation Platform', 3)
    core_deliverables = [
        'Interactive web-based user interface where users can input design specifications (prompts) such as: "Design a robotic bricklaying system for 3-story buildings" or "Create an automated concrete mixing system with temperature control"',
        'RAG (Retrieval-Augmented Generation) system architecture built using Langchain framework that retrieves relevant patent documents from the knowledge base based on user queries',
        'Vector database (e.g., ChromaDB, Pinecone, or FAISS) containing embedded patent documents from construction robotics domain',
        'Semantic search and retrieval system that uses embedding models (e.g., OpenAI embeddings, Sentence-BERT, or custom fine-tuned models) to find most relevant patent documents matching user design specifications',
        'Large Language Model (LLM) integration using Langchain - supports multiple LLM backends including OpenAI GPT-4, GPT-3.5, Anthropic Claude, open-source models (Llama 2, Mistral), or fine-tuned domain-specific models',
        'Prompt engineering and template system using Langchain PromptTemplate and Chain functionality to format retrieved patent context with user design specifications',
        'Automatic design generation module that synthesizes information from retrieved patents and user requirements to generate complete design specifications, technical drawings descriptions, component lists, and implementation guidelines',
        'Design output interface that presents generated designs in structured formats including: technical specifications, component descriptions, workflow diagrams, performance metrics, and design recommendations'
    ]
    for deliverable in core_deliverables:
        doc.add_paragraph(deliverable, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_heading('10.3.2 Patent Knowledge Base and Data Infrastructure', 3)
    knowledge_base_deliverables = [
        'Comprehensive patent document database containing construction robotics patents extracted from USPTO, EPO, WIPO, and Google Patents',
        'Patent document preprocessing pipeline that extracts and structures: titles, abstracts, claims, descriptions, figures, technical specifications, component lists, and performance metrics',
        'Embedded patent document repository with vector representations stored in vector database for efficient semantic search',
        'Patent metadata database including: patent numbers, filing dates, inventors, assignees, classification codes (CPC codes for construction robotics: B25J, E04G, E04B, B66C, etc.)',
        'Knowledge graph construction linking related patents, concepts, components, and design patterns extracted from patent documents',
        'Ontology for construction robotics domain defining relationships between concepts, components, systems, and design patterns'
    ]
    for deliverable in knowledge_base_deliverables:
        doc.add_paragraph(deliverable, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_heading('10.3.3 Langchain and RAG Implementation Components', 3)
    langchain_deliverables = [
        'Langchain-based RAG chain implementation using LCEL (Langchain Expression Language) for composing retrieval and generation workflows',
        'Document loaders and processors using Langchain document loaders for patent document ingestion (PDF, HTML, XML formats)',
        'Text splitters using Langchain RecursiveCharacterTextSplitter or semantic splitters for chunking patent documents into manageable segments',
        'Embedding integrations using Langchain embedding interfaces supporting OpenAI Embeddings, HuggingFace embeddings, Cohere embeddings, and custom embedding models',
        'Vector store integrations using Langchain vector store abstractions for ChromaDB, Pinecone, FAISS, Weaviate, or Qdrant',
        'Retrievers implementation using Langchain retrievers including similarity search, MMR (Maximum Marginal Relevance), and contextual compression retrievers',
        'LLM chains and agents using Langchain LLMChain, SequentialChain, and Agent frameworks for orchestrating design generation workflows',
        'Memory management using Langchain ConversationBufferMemory, ConversationSummaryMemory, or ConversationBufferWindowMemory for maintaining context across design generation sessions',
        'Output parsers using Langchain output parsers to structure generated designs into JSON, XML, or custom formats',
        'Callback handlers and logging using Langchain callbacks for monitoring, debugging, and tracking RAG system performance'
    ]
    for deliverable in langchain_deliverables:
        doc.add_paragraph(deliverable, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_heading('10.3.4 LLM Models and Fine-tuning', 3)
    llm_deliverables = [
        'Fine-tuned LLM models specifically trained on construction robotics patent documents and technical specifications',
        'Model training pipeline and scripts for fine-tuning open-source LLMs (Llama 2, Mistral, CodeLlama) on patent document datasets',
        'Pre-trained model checkpoints and weights for domain-specific construction robotics design generation',
        'Model evaluation framework with benchmarks for assessing design generation quality, relevance, and technical accuracy',
        'Multiple LLM backend support allowing users to choose between different models (GPT-4 for high quality, GPT-3.5 for speed, open-source models for privacy)',
        'Custom prompt templates optimized for construction robotics design generation tasks',
        'Model serving infrastructure for deploying fine-tuned models using Langchain LLM interfaces'
    ]
    for deliverable in llm_deliverables:
        doc.add_paragraph(deliverable, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_heading('10.3.5 System Architecture and Deployment', 3)
    architecture_deliverables = [
        'Complete system architecture documentation including: RAG pipeline flow, component interactions, data flow diagrams, and system deployment architecture',
        'RESTful API endpoints for design generation service allowing integration with external applications',
        'Web application with user-friendly interface for inputting design specifications and viewing generated designs',
        'Database schemas and data models for storing patent documents, embeddings, user queries, and generated designs',
        'Containerization using Docker for easy deployment and scaling of RAG system components',
        'Cloud deployment configurations for AWS, Azure, or GCP including infrastructure as code (Terraform, CloudFormation)',
        'System monitoring and logging infrastructure for tracking system performance, query patterns, and user interactions',
        'Scalability solutions for handling large-scale patent databases and concurrent user requests'
    ]
    for deliverable in architecture_deliverables:
        doc.add_paragraph(deliverable, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_heading('10.3.6 Documentation and Supporting Materials', 3)
    documentation_deliverables = [
        'Comprehensive research paper detailing methodology, RAG architecture, Langchain implementation, experimental results, and findings',
        'Technical documentation including: system architecture, API documentation, code documentation, installation guides, and configuration manuals',
        'User manual and tutorial guides for using the design generation system, including example prompts and design specifications',
        'Developer documentation for extending and customizing the RAG system, including Langchain integration guides and custom chain development',
        'Research dataset repository containing processed patent documents, embeddings, evaluation benchmarks, and case studies',
        'Open-source code repository (GitHub) with complete source code, version control, issue tracking, and contribution guidelines',
        'Performance evaluation reports including: retrieval accuracy metrics, design generation quality assessments, response time benchmarks, and system scalability tests',
        'Case studies demonstrating practical applications of the system with real-world design specifications and generated designs',
        'Video tutorials and demonstration recordings showing system functionality and usage examples',
        'Presentation materials and research posters for academic conferences and industry presentations'
    ]
    for deliverable in documentation_deliverables:
        doc.add_paragraph(deliverable, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The ultimate goal of this project is to deliver a production-ready RAG-based system where users can input '
        'design specifications as natural language prompts (e.g., "Design an automated welding robot for steel '
        'construction") and receive comprehensive design solutions generated by the LLM using relevant patent documents '
        'retrieved through the RAG pipeline. The system leverages Langchain for orchestrating the RAG workflow, '
        'vector databases for efficient patent retrieval, and advanced LLMs for generating high-quality design '
        'specifications based on retrieved knowledge.'
    )
    
    # 11. References
    doc.add_heading('11. References', 1)
    
    references = [
        'Devlin, J., et al. (2018). BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. arXiv preprint arXiv:1810.04805.',
        'Brown, T., et al. (2020). Language Models are Few-Shot Learners. Advances in Neural Information Processing Systems, 33.',
        'Han, J., Kamber, M., & Pei, J. (2011). Data Mining: Concepts and Techniques. Morgan Kaufmann.',
        'Mitchell, T. M. (1997). Machine Learning. McGraw-Hill Education.',
        'Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.',
        'Aggarwal, C. C. (2015). Data Mining: The Textbook. Springer.',
        'Zhu, J., et al. (2019). Patent Mining: A Survey. ACM Computing Surveys, 52(2).',
        'Wang, K., et al. (2020). Automatic Design Generation for Robotics: A Survey. IEEE Transactions on Robotics.'
    ]
    
    for ref in references:
        doc.add_paragraph(ref, style='List Bullet')
    
    # Appendix
    doc.add_page_break()
    doc.add_heading('Appendix A: Key Terminology', 1)
    
    terms = {
        'LLM (Large Language Model)': 'Advanced AI models capable of understanding and generating human-like text by processing vast amounts of linguistic data.',
        'Patent Document': 'Legal documents that describe inventions and grant exclusive rights, containing rich technical information.',
        'Data Mining': 'The process of discovering patterns and knowledge from large amounts of data.',
        'Pattern Recognition': 'The identification of recurring patterns or regularities in data.',
        'Design Specificity': 'Unique characteristics, constraints, or requirements specific to particular design contexts.',
        'Pattern Database': 'A structured repository storing identified patterns, relationships, and rules extracted from data analysis.',
        'Automatic Design': 'The use of computational methods to generate designs automatically based on constraints and objectives.',
        'Knowledge Graph': 'A knowledge representation that uses graph structures to represent entities and relationships.'
    }
    
    for term, definition in terms.items():
        p = doc.add_paragraph()
        p.add_run(f'{term}: ').bold = True
        p.add_run(definition)
    
    doc.add_heading('Appendix B: Research Questions', 1)
    
    research_questions = [
        'How can LLMs be effectively fine-tuned to extract design knowledge from patent documents?',
        'What patterns can be identified in patent documents related to construction robotics technologies?',
        'How can both textual and numeric data be effectively processed and integrated for design applications?',
        'What are the key design specificities that should be extracted from patent documents?',
        'How can pattern databases be structured to support automatic design generation?',
        'What are the optimal data mining approaches for extracting meaningful patterns from construction robotics patents?',
        'How can automatic design systems leverage learned patterns to generate optimal construction robotics designs?'
    ]
    
    for i, question in enumerate(research_questions, 1):
        p = doc.add_paragraph(f'RQ{i}: {question}')
    
    doc.add_page_break()
    doc.add_heading('Appendix C: Patent Database Resources & Links', 1)
    
    doc.add_heading('C.1 Primary Patent Databases', 2)
    
    databases = [
        {
            'name': 'United States Patent and Trademark Office (USPTO)',
            'main_url': 'https://patft.uspto.gov/',
            'advanced': 'https://patft.uspto.gov/netahtml/PTO/search-adv.htm',
            'bulk_data': 'https://bulkdata.uspto.gov/',
            'description': 'Official US patent database with full-text search. Contains all US patents and applications. Bulk data available for research.'
        },
        {
            'name': 'European Patent Office (EPO) - Espacenet',
            'main_url': 'https://worldwide.espacenet.com/',
            'advanced': 'https://worldwide.espacenet.com/advancedSearch?locale=en_EP',
            'bulk_data': 'https://www.epo.org/searching-for-patents/data/bulk-data-sets.html',
            'description': 'Free access to over 130 million patent documents worldwide. Includes European and international patents.'
        },
        {
            'name': 'World Intellectual Property Organization (WIPO) - PATENTSCOPE',
            'main_url': 'https://patentscope.wipo.int/search/en/search.jsf',
            'advanced': 'https://patentscope.wipo.int/search/en/advancedSearch.jsf',
            'bulk_data': 'N/A',
            'description': 'Free access to international patent applications (PCT) and national patent collections from multiple countries.'
        },
        {
            'name': 'Google Patents',
            'main_url': 'https://patents.google.com/',
            'advanced': 'https://patents.google.com/advanced',
            'bulk_data': 'https://console.cloud.google.com/marketplace/product/google_patents_public_datasets',
            'description': 'Comprehensive patent search engine covering multiple patent offices. Includes citation analysis and patent families.'
        },
        {
            'name': 'IEEE Xplore',
            'main_url': 'https://ieeexplore.ieee.org/',
            'advanced': 'N/A',
            'bulk_data': 'N/A',
            'description': 'Technical papers, standards, and patents related to engineering and technology, including robotics.'
        }
    ]
    
    for db in databases:
        p = doc.add_paragraph()
        p.add_run(f'{db["name"]}: ').bold = True
        doc.add_paragraph(f'Main URL: {db["main_url"]}')
        if db['advanced'] != 'N/A':
            doc.add_paragraph(f'Advanced Search: {db["advanced"]}')
        if db['bulk_data'] != 'N/A':
            doc.add_paragraph(f'Bulk Data: {db["bulk_data"]}')
        doc.add_paragraph(db['description'])
        doc.add_paragraph()
    
    doc.add_heading('C.2 Key Patent Classification Codes for Construction Robotics', 2)
    
    classifications = [
        'B25J - Manipulators; Chambers Provided with Manipulation Devices',
        '  • B25J9/00 - Programme-controlled manipulators',
        '  • B25J15/00 - Gripping heads',
        '',
        'E04G - Scaffolding; Forms; Shuttering',
        '  • E04G21/00 - Preparing, conveying, or working-up building materials',
        '  • E04G23/00 - Repairing, securing, or restoring buildings',
        '',
        'E04B - General Building Constructions',
        '  • E04B1/00 - Constructions in general',
        '',
        'B66C - Cranes; Load-Engaging Elements',
        '  • B66C13/00 - Other constructional features or details',
        '',
        'G05B - Control Systems',
        '  • G05B19/00 - Programme-control systems'
    ]
    
    for cls in classifications:
        doc.add_paragraph(cls)
    
    doc.add_paragraph()
    
    doc.add_heading('C.3 Recommended Search Keywords', 2)
    
    keywords = [
        'construction robot',
        'automated construction',
        'robotic construction',
        'construction automation',
        'building robot',
        'construction machinery automation',
        '3D printing construction',
        'prefabrication robot',
        'bricklaying robot',
        'welding robot construction',
        'construction robotics',
        'automated bricklaying',
        'robotic assembly construction'
    ]
    
    for keyword in keywords:
        doc.add_paragraph(keyword, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('C.4 Data Access Methods', 2)
    
    doc.add_paragraph('1. Web Interface: Direct search through database websites')
    doc.add_paragraph('2. Bulk Downloads: Download patent data in bulk formats (XML, JSON)')
    doc.add_paragraph('3. APIs: Use official APIs where available (USPTO, EPO)')
    doc.add_paragraph('4. Academic Access: University library subscriptions to patent databases')
    
    doc.add_paragraph()
    
    doc.add_heading('C.5 Additional Resources', 2)
    
    resources = [
        'USPTO Patent Classification: https://www.uspto.gov/web/patents/classification/',
        'CPC (Cooperative Patent Classification): https://www.cooperativepatentclassification.org/',
        'USPTO Developer Portal: https://developer.uspto.gov/',
        'EPO Open Patent Services: https://developers.epo.org/',
        'Construction Robotics Research: International Association for Automation and Robotics in Construction (IAARC)'
    ]
    
    for resource in resources:
        doc.add_paragraph(resource, style='List Bullet')
    
    # Save document
    output_path = '/Users/anuj_kittur/Desktop/Scraper/Research_Proposal_Patent_LLM_Design.docx'
    doc.save(output_path)
    print(f"Document 2 created successfully: {output_path}")
    
    return output_path

if __name__ == '__main__':
    create_patent_llm_document()

